# -*- coding: utf-8 -*-
"""사용설명서.md → 사용설명서.docx (python-docx). 제목·표·목록·인라인 코드/굵게를 옮긴다.
실행: python docs/ksx3267/manual/md2docx.py
"""
import io
import os
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, "사용설명서.md")
OUT = os.path.join(HERE, "사용설명서.docx")
FONT = "맑은 고딕"


def set_font(run, size=None, bold=None, mono=False, color=None):
    run.font.name = "Consolas" if mono else FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas" if mono else FONT)
    if size: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if color: run.font.color.rgb = RGBColor(*color)


def shade(cell, hex_fill):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


INLINE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


def add_inline(par, text, size=10.5):
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            set_font(par.add_run(part[2:-2]), size=size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            set_font(par.add_run(part[1:-1]), size=size - 0.5, mono=True, color=(0x1F, 0x3A, 0x5F))
        else:
            set_font(par.add_run(part), size=size)


def add_table(doc, rows):
    header, body = rows[0], [r for r in rows[1:] if not re.match(r"^\s*\|?\s*-{2,}", r)]
    split = lambda r: [c.strip() for c in r.strip().strip("|").split("|")]
    cols = split(header)
    t = doc.add_table(rows=1, cols=len(cols))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, c in enumerate(cols):
        cell = t.rows[0].cells[i]; cell.text = ""
        add_inline(cell.paragraphs[0], c, size=9.5); cell.paragraphs[0].runs[0].bold = True
        shade(cell, "E8EEF7")
    for r in body:
        cells = split(r)
        row = t.add_row().cells
        for i in range(len(cols)):
            row[i].text = ""
            add_inline(row[i].paragraphs[0], cells[i] if i < len(cells) else "", size=9.5)
    doc.add_paragraph()


def build():
    doc = Document()
    st = doc.styles["Normal"]; st.font.name = FONT; st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT); st.font.size = Pt(10.5)
    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(2.2); s.top_margin = s.bottom_margin = Cm(2.0)

    lines = io.open(SRC, encoding="utf-8").read().splitlines()
    i = 0
    first_h1 = True
    while i < len(lines):
        ln = lines[i]
        if ln.startswith("# "):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(p.add_run(ln[2:].strip()), size=20, bold=True, color=(0x1F, 0x3A, 0x5F))
            if first_h1:
                first_h1 = False
                # 다음 두 줄(버전·대상)을 부제로
                j = i + 1
                while j < len(lines) and lines[j].strip():
                    q = doc.add_paragraph(); q.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    add_inline(q, lines[j].strip(), size=10)
                    j += 1
                i = j; continue
        elif ln.startswith("## "):
            h = doc.add_heading(level=1); set_font(h.add_run(ln[3:].strip()), size=14, bold=True, color=(0x1F, 0x3A, 0x5F))
        elif ln.startswith("### "):
            h = doc.add_heading(level=2); set_font(h.add_run(ln[4:].strip()), size=12, bold=True, color=(0x2E, 0x4A, 0x6E))
        elif ln.strip().startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i]); i += 1
            add_table(doc, block); continue
        elif re.match(r"^\s*\d+\.\s", ln):
            p = doc.add_paragraph(style="List Number"); add_inline(p, re.sub(r"^\s*\d+\.\s", "", ln))
        elif re.match(r"^\s*[-•]\s", ln):
            indent = len(ln) - len(ln.lstrip())
            p = doc.add_paragraph(style="List Bullet 2" if indent >= 2 else "List Bullet"); add_inline(p, re.sub(r"^\s*[-•]\s", "", ln))
        elif ln.strip():
            p = doc.add_paragraph(); add_inline(p, ln.strip())
        i += 1

    # 바닥글: 문서 식별
    footer = doc.sections[0].footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("스마트그린 통합제어기 — KS X 3267 표준 노드 사용설명서 v1.0"), size=8, color=(0x80, 0x80, 0x80))
    doc.save(OUT)
    print("wrote", OUT)


if __name__ == "__main__":
    build()
